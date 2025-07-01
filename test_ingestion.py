import os
import pytest
from unittest.mock import patch, MagicMock
from pathlib import Path
import ingest  # this assumes your script is named ingest.py

@pytest.fixture(scope="module")
def sample_docx(tmp_path_factory):
    # Create a dummy DOCX file
    from docx import Document
    tmp_path = tmp_path_factory.mktemp("data")
    doc_path = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("Test paragraph for unit testing.")
    doc.save(doc_path)
    return doc_path

def test_extract_text_from_docx(sample_docx):
    text = ingest.extract_text_from_docx(sample_docx)
    assert "Test paragraph for unit testing." in text

def test_load_documents_from_directory(tmp_path):
    # Create a dummy DOCX file
    from docx import Document
    doc_path = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("Sample paragraph.")
    doc.save(doc_path)

    docs = ingest.load_documents_from_directory(tmp_path)
    assert len(docs) == 1
    assert docs[0]["type"] == "docx"

def test_get_document_text_from_docx():
    documents = [{
        "name": "test.docx",
        "type": "docx",
        "content": "This is test content from a DOCX file."
    }]
    text_by_source = ingest.get_document_text(documents)
    assert "test.docx" in text_by_source
    assert "This is test content" in text_by_source["test.docx"]

def test_get_text_chunks():
    text_by_source = {
        "sample.txt": "Line 1.\nLine 2.\n" * 100
    }
    chunks, metadata = ingest.get_text_chunks(text_by_source)
    assert len(chunks) > 0
    assert len(chunks) == len(metadata)
    assert "sample.txt" in metadata[0]["source"]

@patch("ingest.get_vectorstore", return_value=True)
@patch("ingest.get_text_chunks", return_value=(["test chunk"], [{"source": "test.docx"}]))
@patch("ingest.get_document_text", return_value={"test.docx": "text"})
@patch("ingest.load_documents_from_directory", return_value=[{"name": "test.docx", "type": "docx", "content": "text"}])
def test_ingest_documents(mock_load, mock_text, mock_chunks, mock_vectorstore):
    result = ingest.ingest_documents()
    assert result is True
